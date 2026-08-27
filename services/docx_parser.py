"""
DOCX Question Parser — Enhanced parser for DOCX files with OMML & MathType equations.

Strategy:
  1. Use `python-docx` to parse question structure, OMML math → LaTeX,
     options, statements, and explanations.
  2. MathType WMF/OLE objects are identified, raw MTEF is extracted, and
     SHA-256 is computed on raw bytes before zlib+base64 compression.
  3. Embedded TeX/MathML metadata from MathType translators is used if present.
  4. MTEF without metadata stays `pending` for external worker conversion.
  5. OMML pipeline is completely separate from MathType.
  6. Emit `[[formula:<uuid>]]` placeholders and collect formulas per paragraph.
"""

import base64
import hashlib
import os
import re
import shutil
import subprocess
import tempfile
import uuid
import zlib
import xml.etree.ElementTree as ET

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

UPLOAD_QUESTIONS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static", "uploads", "questions")

def ensure_questions_dir():
    os.makedirs(UPLOAD_QUESTIONS_DIR, exist_ok=True)


def _crop_libreoffice_svg(svg_path):
    """Crop LibreOffice's A4 Draw export to the WMF graphic bounds."""
    try:
        with open(svg_path, "r", encoding="utf-8") as svg_file:
            content = svg_file.read()
        bbox = re.search(
            r'<rect class="BoundingBox"[^>]*\bx="([\d.]+)"[^>]*\by="([\d.]+)"'
            r'[^>]*\bwidth="([\d.]+)"[^>]*\bheight="([\d.]+)"',
            content,
        )
        if not bbox:
            return
        x, y, width, height = (float(value) for value in bbox.groups())
        padding = max(40.0, height * 0.06)
        x = max(0.0, x - padding)
        y = max(0.0, y - padding)
        width += padding * 2
        height += padding * 2
        replacement = (
            f'<svg version="1.2" width="{width / 100:.2f}mm" '
            f'height="{height / 100:.2f}mm" '
            f'viewBox="{x:.0f} {y:.0f} {width:.0f} {height:.0f}"'
        )
        content = re.sub(
            r'<svg version="1\.2"\s+width="[^"]+"\s+height="[^"]+"\s+'
            r'viewBox="[^"]+"',
            replacement,
            content,
            count=1,
        )
        with open(svg_path, "w", encoding="utf-8") as svg_file:
            svg_file.write(content)
    except (OSError, ValueError):
        return


def _convert_wmf_previews_to_svg(previews):
    """Batch-convert Word's MathType WMF previews to browser SVG files.

    The SVG is a derived display cache, not the canonical formula payload.  If
    LibreOffice is unavailable the import still succeeds and the MTEF payload
    remains available for the configured conversion worker.

    Args:
        previews: mapping of relationship ID to WMF bytes.
    Returns:
        mapping of relationship ID to web-accessible SVG URL.
    """
    soffice = shutil.which("soffice")
    if not soffice or not previews:
        return {}

    ensure_questions_dir()
    task_dir = tempfile.mkdtemp(prefix="mathtype_svg_")
    profile_dir = tempfile.mkdtemp(prefix="mathtype_lo_profile_")
    source_paths = []
    source_names = {}
    converted = {}

    try:
        for index, (rel_id, wmf_bytes) in enumerate(previews.items()):
            stem = f"formula_{index}_{uuid.uuid4().hex[:8]}"
            source_path = os.path.join(task_dir, stem + ".wmf")
            with open(source_path, "wb") as source_file:
                source_file.write(wmf_bytes)
            source_paths.append(source_path)
            source_names[rel_id] = stem

        profile_uri = "file://" + profile_dir
        result = subprocess.run(
            [
                soffice,
                f"-env:UserInstallation={profile_uri}",
                "--headless",
                "--convert-to",
                "svg",
                "--outdir",
                task_dir,
            ] + source_paths,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=max(30, len(source_paths) * 3),
            check=False,
        )
        if result.returncode != 0:
            return {}

        for rel_id, stem in source_names.items():
            output_path = os.path.join(task_dir, stem + ".svg")
            if not os.path.exists(output_path):
                continue
            final_name = f"formula_{uuid.uuid4().hex[:12]}.svg"
            final_path = os.path.join(UPLOAD_QUESTIONS_DIR, final_name)
            shutil.copyfile(output_path, final_path)
            _crop_libreoffice_svg(final_path)
            converted[rel_id] = f"/static/uploads/questions/{final_name}"
        return converted
    except (OSError, subprocess.SubprocessError):
        return {}
    finally:
        shutil.rmtree(task_dir, ignore_errors=True)
        shutil.rmtree(profile_dir, ignore_errors=True)


def _relationship_blob(rel):
    """Return bytes for an internal DOCX relationship.

    ``python-docx`` normally exposes ``target_part.blob``.  Older Word files
    and a few MathType exporters can leave a relationship whose target is
    still available through the private ``_target`` attribute, while
    ``target_part`` raises or is not present.  The fallback is deliberately
    read-only and keeps the parser compatible with both forms.
    """
    try:
        target_part = getattr(rel, "target_part", None)
        blob = getattr(target_part, "blob", None)
        if blob:
            return blob
    except Exception:
        pass

    try:
        target = getattr(rel, "_target", None)
        blob = getattr(target, "blob", None)
        if blob:
            return blob
    except Exception:
        pass
    return None


from services.mathtype_converter import process_mathtype_formula

# ---------------------------------------------------------------------------
#  MathType MTEF to LaTeX extraction
# ---------------------------------------------------------------------------
# Delegated to services.mathtype_converter


# ---------------------------------------------------------------------------
#  Main DOCX parser
# ---------------------------------------------------------------------------

def parse_docx(file_stream_or_path):
    """
    Parse DOCX file.
    Returns: list of raw question dicts (with 'formulas' dictionary containing {uuid: {latex, ...}})
    """
    if not HAS_DOCX:
        raise RuntimeError("Thư viện 'python-docx' chưa được cài đặt.")

    if hasattr(file_stream_or_path, "read"):
        file_stream_or_path.seek(0)
        doc = docx.Document(file_stream_or_path)
    else:
        doc = docx.Document(file_stream_or_path)

    ensure_questions_dir()

    # ------- Phase 1: Extract standard images and MathType formulas -------
    rel_id_to_image = {}
    wmf_preview_sources = {}
    stats = {"standard_imgs": 0, "wmf_extracted": 0, "wmf_skipped": 0}

    try:
        for rel_id, rel in doc.part.rels.items():
            target_part = getattr(rel, "target_part", None)
            content_type = getattr(target_part, "content_type", "").lower()
            rel_type = getattr(rel, "reltype", "").lower()
            target_ref = str(getattr(rel, "target_ref", "")).lower()
            is_ole = (
                "ole" in content_type
                or "oleobject" in rel_type
                or "oleobject" in target_ref
                or target_ref.endswith(".bin")
            )
            is_image = "image" in content_type

            # Some python-docx versions cannot resolve target_part for an
            # embedded OLE relationship, although the relationship itself is
            # valid.  Detect it from the relationship metadata and read the
            # target through the compatibility helper above.
            if not target_part and not (is_ole or is_image):
                continue

            img_bytes = _relationship_blob(rel)
            if not img_bytes:
                if is_ole:
                    stats["wmf_skipped"] += 1
                continue

            if is_image or is_ole:
                ct = content_type.split("/")[-1]

                if is_ole or ct in ("x-wmf", "wmf", "x-emf", "emf", "oleobject"):
                    provider_result = process_mathtype_formula(img_bytes)
                    
                    if provider_result.get("mtef_base64") or provider_result.get("latex") or provider_result.get("mathml"):
                        f_id = str(uuid.uuid4())
                        
                        # Determine review status from conversion pipeline
                        conversion_status = provider_result.get("conversion_status", "pending")
                        needs_review = conversion_status == "pending" or provider_result.get("needs_review", True)
                        
                        # Build formula data dict with production fields
                        formula_data = {
                            "latex": provider_result.get("latex"),
                            "mathml": provider_result.get("mathml"),
                            "mtef_data": provider_result.get("mtef_base64"),
                            "content_hash": provider_result.get("content_hash"),
                            "source_format": "MathType",
                            "needs_review": needs_review,
                            "conversion_status": conversion_status,
                            "converter_name": provider_result.get("provider_name"),
                            "parse_confidence": provider_result.get("confidence", 0.0),
                        }

                        # MathType stores a faithful WMF preview beside/in the
                        # OLE object.  It gives an immediate, automatic display
                        # fallback while structured MTEF conversion is pending.
                        if ct in ("x-wmf", "wmf"):
                            wmf_preview_sources[rel_id] = img_bytes
                        
                        rel_id_to_image[rel_id] = {
                            "id": f_id,
                            "data": formula_data,
                        }
                        stats["wmf_extracted"] += 1
                    else:
                        stats["wmf_skipped"] += 1
                    continue
                
                # Standard Images
                if ct in ("jpeg", "pjpeg"): ext = "jpg"
                elif ct in ("png",): ext = "png"
                elif ct in ("gif",): ext = "gif"
                else: ext = ct if ct in ("bmp", "tiff", "svg+xml") else "png"

                filename = f"img_docx_{uuid.uuid4().hex[:8]}.{ext}"
                filepath = os.path.join(UPLOAD_QUESTIONS_DIR, filename)
                with open(filepath, "wb") as f:
                    f.write(img_bytes)
                web_path = f"/static/uploads/questions/{filename}"
                rel_id_to_image[rel_id] = web_path
                stats["standard_imgs"] += 1
    except Exception as e:
        print(f"[docx_parser] Error mapping docx images: {e}")

    # Convert all MathType previews in one LibreOffice process. Starting one
    # office process per formula makes medium-sized question banks appear hung.
    for rel_id, preview_url in _convert_wmf_previews_to_svg(wmf_preview_sources).items():
        mapped = rel_id_to_image.get(rel_id)
        if isinstance(mapped, dict):
            mapped["data"]["preview_url"] = preview_url

    print(f"[docx_parser] Image stats: {stats}")

    # ------- Phase 2: Parse question structure from paragraphs -------
    raw_questions = []
    current_q = None
    in_explanation = False

    q_num_pattern = re.compile(
        r'^(?:<[^>]+>\s*)*(?:\[|\(|\【)?\s*(?:CÂU|Câu|Question|Bài|Ví dụ|Q|Q\.|Phần|Item)\s*(\d+)\s*[.\:)\-】\]]*\s*(.*)',
        re.IGNORECASE
    )
    alt_num_pattern = re.compile(r'^(?:\[|\()?(\d+)[\.\:\)\-]\s+(.*)')
    option_pattern = re.compile(r"^(['’'*]?)\s*(?:\[|\()?([A-E])[.:)\-\]\】]\s*(.*)")
    inline_option_pattern = re.compile(r"(?:^|\s+|\t|(?<=[.,;!?:$]))(['’'*]?)\s*(?:\[|\()?([A-E])[.:)\-\]\】]\s*")
    tf_stmt_pattern = re.compile(r'^([a-d])[.:)]\s*(.*)')
    type_pattern = re.compile(r'^(?:TYPE|LOẠI|Loại câu hỏi)[:\s]+(.*)', re.IGNORECASE)
    answer_pattern = re.compile(r'^(?:ANSWER|Đáp án|ĐÁP ÁN|Đáp án đúng|Chọn)[:\s]+(.*)', re.IGNORECASE)
    diff_pattern = re.compile(r'^(?:Mức độ|MỨC ĐỘ|Level)[:\s]+(.*)', re.IGNORECASE)
    exp_pattern = re.compile(r'^(?:Giải thích|GIẢI THÍCH|Lời giải|Lời giải:|Hướng dẫn giải|HDG|Lời giải tham khảo)[:\s]*(.*)', re.IGNORECASE)

    all_paragraphs = list(_iter_all_paragraphs(doc))
    unassigned_text_lines = []
    unassigned_formulas = {}

    for p_idx, p in enumerate(all_paragraphs):
        p_formulas = {}
        p_text = _extract_paragraph_text(p, rel_id_to_image, p_formulas).strip()
        q_match = q_num_pattern.match(p_text) or alt_num_pattern.match(p_text)
        if q_match:
            if current_q:
                raw_questions.append(current_q)
            in_explanation = False
            q_text = q_match.group(2).strip()
            
            new_formulas = dict(unassigned_formulas)
            # A question heading can contain its first MathType object on
            # the same paragraph ("Câu 1. Tính ... <MathType>").  Attach
            # those formulas to the new question, not to the previous one.
            new_formulas.update(p_formulas)
            if not q_text and unassigned_text_lines:
                q_text = "\n".join(unassigned_text_lines)
            unassigned_text_lines = []
            unassigned_formulas = {}

            current_q = {
                "question_text": q_text,
                "context": "",
                "question_type": "",
                "difficulty_level": "",
                "explanation": "",
                "options": [],
                "statements": [],
                "raw_correct": "",
                "image_url": "",
                "formulas": new_formulas,
                "confidence_scores": {
                    "question": 0.98,
                    "type": 0.95,
                    "image": 1.0,
                    "answer": 0.80,
                },
            }
            continue

        # Formulas in ordinary paragraphs belong to the current question;
        # formulas before the first heading are carried into the next one.
        if current_q:
            current_q["formulas"].update(p_formulas)
        else:
            unassigned_formulas.update(p_formulas)

        if not p_text:
            continue

        inline_matches = list(inline_option_pattern.finditer(p_text))
        if not current_q and inline_matches:
            q_text = "\n".join(unassigned_text_lines).strip() if unassigned_text_lines else f"Câu hỏi {len(raw_questions) + 1}"
            
            new_formulas = dict(unassigned_formulas)
            unassigned_text_lines = []
            unassigned_formulas = {}
            
            current_q = {
                "question_text": q_text,
                "context": "",
                "question_type": "",
                "difficulty_level": "",
                "explanation": "",
                "options": [],
                "statements": [],
                "raw_correct": "",
                "image_url": "",
                "formulas": new_formulas,
                "confidence_scores": {
                    "question": 0.90,
                    "type": 0.90,
                    "image": 1.0,
                    "answer": 0.80,
                },
            }

        if not current_q:
            unassigned_text_lines.append(p_text)
            continue

        t_match = type_pattern.match(p_text)
        if t_match:
            current_q["question_type"] = _map_template_type(t_match.group(1).strip())
            continue

        exp_match = exp_pattern.match(p_text)
        if exp_match and not in_explanation:
            in_explanation = True
            first_exp = exp_match.group(1).strip()
            if first_exp:
                current_q["explanation"] = first_exp
            continue

        if in_explanation:
            if current_q["explanation"]:
                current_q["explanation"] += "\n" + p_text
            else:
                current_q["explanation"] = p_text
            continue

        inline_matches = list(inline_option_pattern.finditer(p_text))
        if inline_matches and inline_matches[0].start() <= 3:
            for idx_m, m in enumerate(inline_matches):
                prefix = m.group(1)
                letter = m.group(2).upper()
                start_idx = m.end()
                end_idx = inline_matches[idx_m + 1].start() if idx_m + 1 < len(inline_matches) else len(p_text)
                
                opt_text = p_text[start_idx:end_idx].strip()
                is_corr = bool(prefix in ("'", "’", "’", "*")) or bool(
                    re.search(r'[*]|(?:\(đúng\))|(?:\(chính xác\))$', opt_text, re.IGNORECASE)
                )
                if opt_text.endswith("*"):
                    opt_text = opt_text[:-1].strip()

                current_q["options"].append({
                    "letter": letter,
                    "text": opt_text,
                    "is_correct": is_corr,
                    "order_index": len(current_q["options"]),
                })
            continue

        stmt_match = tf_stmt_pattern.match(p_text)
        if stmt_match and len(current_q["options"]) == 0:
            stmt_id = stmt_match.group(1).lower()
            stmt_text = stmt_match.group(2).strip()
            ans_val = None
            if re.search(r'\b(là đúng|ĐÚNG|\(Đ\)|TRUE)\b', stmt_text, re.IGNORECASE):
                ans_val = True
            elif re.search(r'\b(là sai|SAI|\(S\)|FALSE)\b', stmt_text, re.IGNORECASE):
                ans_val = False

            current_q["statements"].append({
                "id": stmt_id,
                "text": stmt_text,
                "answer": ans_val,
            })
            current_q["question_type"] = "true_false"
            continue

        ans_match = answer_pattern.match(p_text)
        if ans_match:
            current_q["raw_correct"] = ans_match.group(1).strip()
            continue

        diff_match = diff_pattern.match(p_text)
        if diff_match:
            current_q["difficulty_level"] = diff_match.group(1).strip()
            continue

        if not current_q["options"] and not current_q["statements"] and not current_q["raw_correct"]:
            current_q["question_text"] += "\n" + p_text

    if current_q:
        raw_questions.append(current_q)

    for q in raw_questions:
        raw_corr = q.get("raw_correct", "").upper()
        if raw_corr and q["options"]:
            corr_letters = set(re.findall(r"[A-E]", raw_corr))
            for opt in q["options"]:
                if opt.get("letter") in corr_letters or opt.get("text", "").upper() == raw_corr:
                    opt["is_correct"] = True

    return raw_questions


# ---------------------------------------------------------------------------
#  Paragraph / XML walking helpers
# ---------------------------------------------------------------------------

def _iter_all_paragraphs(doc):
    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            yield docx.text.paragraph.Paragraph(child, doc)
        elif tag == "tbl":
            table = docx.table.Table(child, doc)
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p


def _extract_paragraph_text(p, rel_id_to_image, formula_dict):
    p_xml = p._element
    chunks = []
    _walk_xml_node(p_xml, rel_id_to_image, chunks, formula_dict)
    full_text = "".join(chunks).strip()
    return full_text if full_text else (p.text or "")


def _walk_xml_node(elem, rel_id_to_image, chunks, formula_dict):
    tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag

    if tag == "object":
        # A MathType object normally contains two relationships: a WMF/PNG
        # preview and the actual OLE object.  The preview is often visited
        # first in the XML.  Always prefer the formula relation so the image
        # preview cannot hide the MTEF payload.
        relation_ids = []
        for child_elem in elem.iter():
            for attr_key, attr_val in child_elem.attrib.items():
                attr_name = attr_key.split("}")[-1] if "}" in attr_key else attr_key
                if attr_name in ("embed", "id", "href") and attr_val in rel_id_to_image:
                    if attr_val not in relation_ids:
                        relation_ids.append(attr_val)

        formula_ids = [rid for rid in relation_ids if isinstance(rel_id_to_image[rid], dict)]
        preview_ids = [rid for rid in relation_ids if not isinstance(rel_id_to_image[rid], dict)]
        for rel_id in formula_ids + [rid for rid in relation_ids if rid not in formula_ids]:
            img_val = rel_id_to_image[rel_id]
            if isinstance(img_val, dict):
                # Keep the exact MathType-rendered preview available while
                # an optional MathType worker converts the MTEF payload.
                if preview_ids:
                    img_val["data"].setdefault("preview_url", rel_id_to_image[preview_ids[0]])
                f_id = img_val["id"]
                formula_dict[f_id] = img_val["data"]
                chunks.append(f" [[formula:{f_id}]] ")
            else:
                chunks.append(f' <img src="{img_val}" class="docx-math-img" style="vertical-align:middle; max-height:45px; margin:0 4px; display:inline-block;" /> ')
            return
        return

    if tag in ("oMath", "oMathPara"):
        latex_math = _omml_to_latex(elem)
        if latex_math:
            f_id = str(uuid.uuid4())
            formula_dict[f_id] = {
                "latex": latex_math,
                "source_format": "OMML",
                "needs_review": False,
                "content_hash": hashlib.sha256(latex_math.encode("utf-8")).hexdigest(),
                "conversion_status": "converted",
                "parse_confidence": 1.0,
            }
            chunks.append(f" [[formula:{f_id}]] ")
        return

    if tag == "t" and elem.text:
        chunks.append(elem.text)
        return

    found_alt = False
    for attr_key, attr_val in elem.attrib.items():
        attr_name = attr_key.split("}")[-1] if "}" in attr_key else attr_key
        if attr_name in ("alt", "descr", "title") and attr_val:
            val_str = str(attr_val).strip()
            if _is_math_latex_str(val_str):
                clean_str = val_str.strip("$")
                f_id = str(uuid.uuid4())
                formula_dict[f_id] = {
                    "latex": clean_str,
                    "source_format": "LaTeX",
                    "needs_review": False,
                    "content_hash": hashlib.sha256(clean_str.encode("utf-8")).hexdigest(),
                    "conversion_status": "converted",
                    "parse_confidence": 1.0,
                }
                chunks.append(f" [[formula:{f_id}]] ")
                found_alt = True
                break

    if found_alt:
        return

    found_img = False
    for attr_key, attr_val in elem.attrib.items():
        attr_name = attr_key.split("}")[-1] if "}" in attr_key else attr_key
        if attr_name in ("embed", "id", "href") and attr_val in rel_id_to_image:
            img_val = rel_id_to_image[attr_val]
            if isinstance(img_val, dict):
                f_id = img_val["id"]
                formula_dict[f_id] = img_val["data"]
                chunks.append(f" [[formula:{f_id}]] ")
            else:
                chunks.append(f' <img src="{img_val}" class="docx-math-img" style="vertical-align:middle; max-height:45px; margin:0 4px; display:inline-block;" /> ')
            found_img = True
            break

    if found_img:
        return

    for child in elem:
        _walk_xml_node(child, rel_id_to_image, chunks, formula_dict)


# ---------------------------------------------------------------------------
#  Math detection & OMML→LaTeX converter
# ---------------------------------------------------------------------------

def _is_math_latex_str(s):
    if not s or len(str(s).strip()) < 2:
        return False
    s_clean = str(s).strip()
    if s_clean.startswith("$") and s_clean.endswith("$"):
        return True
    indicators = [
        "\\frac", "\\sqrt", "\\sum", "\\int", "\\lim",
        "\\alpha", "\\beta", "\\gamma", "\\delta", "\\theta", "\\pi",
        "\\mathbb", "\\mathcal", "\\forall", "\\exists",
        "^{", "_{",
    ]
    return any(ind in s_clean for ind in indicators)


def _omml_to_latex(elem):
    tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag

    if tag == "t":
        return elem.text or ""

    if tag in ("oMath", "oMathPara", "e", "sub", "sup", "num", "den", "deg"):
        return "".join([_omml_to_latex(c) for c in elem])

    if tag == "r":
        parts = []
        for c in elem:
            ctag = c.tag.split("}")[-1] if "}" in c.tag else c.tag
            if ctag == "t" and c.text:
                parts.append(c.text)
        return "".join(parts)

    if tag == "f":
        num_elem = den_elem = None
        expr_elems = []
        for c in elem:
            ctag = c.tag.split("}")[-1] if "}" in c.tag else c.tag
            if ctag == "num":
                num_elem = c
            elif ctag == "den":
                den_elem = c
            elif ctag == "e":
                # Some producers (and older Word-generated XML) use two
                # direct m:e children instead of m:num/m:den.
                expr_elems.append(c)
        if num_elem is None and expr_elems:
            num_elem = expr_elems[0]
        if den_elem is None and len(expr_elems) > 1:
            den_elem = expr_elems[1]
        num = _omml_to_latex(num_elem) if num_elem is not None else ""
        den = _omml_to_latex(den_elem) if den_elem is not None else ""
        return f"\\frac{{{num}}}{{{den}}}"

    if tag == "rad":
        deg_elem = e_elem = None
        for c in elem:
            ctag = c.tag.split("}")[-1] if "}" in c.tag else c.tag
            if ctag == "deg":
                deg_elem = c
            elif ctag == "e":
                e_elem = c
        expr = _omml_to_latex(e_elem) if e_elem is not None else ""
        deg_text = _omml_to_latex(deg_elem).strip() if deg_elem is not None else ""
        if deg_text and deg_text != "2":
            return f"\\sqrt[{deg_text}]{{{expr}}}"
        return f"\\sqrt{{{expr}}}"

    if tag == "sSup":
        e_elem = sup_elem = None
        for c in elem:
            ctag = c.tag.split("}")[-1] if "}" in c.tag else c.tag
            if ctag == "e":
                e_elem = c
            elif ctag == "sup":
                sup_elem = c
        base = _omml_to_latex(e_elem) if e_elem is not None else ""
        sup = _omml_to_latex(sup_elem) if sup_elem is not None else ""
        return f"{{{base}}}^{{{sup}}}"

    if tag == "sSub":
        e_elem = sub_elem = None
        for c in elem:
            ctag = c.tag.split("}")[-1] if "}" in c.tag else c.tag
            if ctag == "e":
                e_elem = c
            elif ctag == "sub":
                sub_elem = c
        base = _omml_to_latex(e_elem) if e_elem is not None else ""
        sub = _omml_to_latex(sub_elem) if sub_elem is not None else ""
        return f"{{{base}}}_{{{sub}}}"

    if tag == "sSubSup":
        e_elem = sub_elem = sup_elem = None
        for c in elem:
            ctag = c.tag.split("}")[-1] if "}" in c.tag else c.tag
            if ctag == "e":
                e_elem = c
            elif ctag == "sub":
                sub_elem = c
            elif ctag == "sup":
                sup_elem = c
        base = _omml_to_latex(e_elem) if e_elem is not None else ""
        sub = _omml_to_latex(sub_elem) if sub_elem is not None else ""
        sup = _omml_to_latex(sup_elem) if sup_elem is not None else ""
        return f"{{{base}}}_{{{sub}}}^{{{sup}}}"

    if tag == "d":
        inner_parts = []
        for c in elem:
            ctag = c.tag.split("}")[-1] if "}" in c.tag else c.tag
            if ctag == "e":
                inner_parts.append(_omml_to_latex(c))
        content = ",".join(inner_parts)
        return f"\\left({content}\\right)"

    if tag == "nary":
        sub_elem = sup_elem = e_elem = None
        chr_val = ""
        for c in elem:
            ctag = c.tag.split("}")[-1] if "}" in c.tag else c.tag
            if ctag == "naryPr":
                for prop in c:
                    ptag = prop.tag.split("}")[-1] if "}" in prop.tag else prop.tag
                    if ptag == "chr":
                        for ak, av in prop.attrib.items():
                            aname = ak.split("}")[-1] if "}" in ak else ak
                            if aname == "val":
                                chr_val = av
            elif ctag == "sub":
                sub_elem = c
            elif ctag == "sup":
                sup_elem = c
            elif ctag == "e":
                e_elem = c

        op_map = {"∑": "\\sum", "∫": "\\int", "∏": "\\prod", "∐": "\\coprod"}
        op = op_map.get(chr_val, chr_val or "\\sum")
        sub = _omml_to_latex(sub_elem) if sub_elem is not None else ""
        sup = _omml_to_latex(sup_elem) if sup_elem is not None else ""
        expr = _omml_to_latex(e_elem) if e_elem is not None else ""
        result = op
        if sub:
            result += f"_{{{sub}}}"
        if sup:
            result += f"^{{{sup}}}"
        result += f" {expr}"
        return result

    if tag == "limLow":
        e_elem = lim_elem = None
        for c in elem:
            ctag = c.tag.split("}")[-1] if "}" in c.tag else c.tag
            if ctag == "e":
                e_elem = c
            elif ctag == "lim":
                lim_elem = c
        base = _omml_to_latex(e_elem) if e_elem is not None else ""
        lim = _omml_to_latex(lim_elem) if lim_elem is not None else ""
        return f"{base}_{{{lim}}}"

    if tag == "bar":
        e_elem = None
        for c in elem:
            ctag = c.tag.split("}")[-1] if "}" in c.tag else c.tag
            if ctag == "e":
                e_elem = c
        expr = _omml_to_latex(e_elem) if e_elem is not None else ""
        return f"\\overline{{{expr}}}"

    if tag == "m":
        rows = []
        for c in elem:
            ctag = c.tag.split("}")[-1] if "}" in c.tag else c.tag
            if ctag == "mr":
                cells = []
                for cell in c:
                    ctag2 = cell.tag.split("}")[-1] if "}" in cell.tag else cell.tag
                    if ctag2 == "e":
                        cells.append(_omml_to_latex(cell))
                rows.append(" & ".join(cells))
        return "\\begin{matrix}" + " \\\\ ".join(rows) + "\\end{matrix}"

    return "".join([_omml_to_latex(c) for c in elem])


def _map_template_type(raw_type):
    t = raw_type.upper()
    if "SINGLE" in t or "TRẮC NGHIỆM" in t:
        return "single"
    if "MULTIPLE" in t or "NHIỀU ĐÁP ÁN" in t:
        return "multiple_choice"
    if "TRUE" in t or "ĐÚNG" in t:
        return "true_false"
    if "SHORT" in t or "NGẮN" in t:
        return "short_answer"
    if "ESSAY" in t or "TỰ LUẬN" in t:
        return "essay"
    return "single"
